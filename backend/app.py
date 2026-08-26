from flask import Flask, request
from flask_cors import CORS
from pypdf import PdfReader
import docx
import os
import json
from dotenv import load_dotenv
from google import genai

load_dotenv()

app = Flask(__name__)
CORS(app)

client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

MAX_FILE_SIZE_MB = 10
ALLOWED_EXTENSIONS = [".pdf", ".docx"]

with open("pdpa_checklist.json", "r") as f:
    PDPA_CHECKLIST = json.load(f)

def extract_text_from_pdf(file):
    reader = PdfReader(file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file):
    document = docx.Document(file)
    text = ""

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            text += row_text + "\n"

    return text

def build_checklist_summary():
    lines = []
    for item in PDPA_CHECKLIST:
        lines.append(
            f"- [{item['id']}] {item['category']} (Severity: {item['severity']}): "
            f"{item['requirement']} What to look for: {item['what_to_look_for']}"
        )
    return "\n".join(lines)

def build_gap_analysis_prompt(document_text):
    checklist_text = build_checklist_summary()

    prompt = f"""You are a PDPA (Personal Data Protection Act, Sri Lanka) compliance analyst.

Below is a checklist of {len(PDPA_CHECKLIST)} PDPA requirements. For each requirement, review the provided document and determine whether the document shows evidence that the requirement is met, partially met, or not addressed at all.

CHECKLIST:
{checklist_text}

DOCUMENT TO ANALYZE:
{document_text[:8000]}

Respond ONLY with valid JSON (no markdown formatting, no code blocks, no extra text) in this exact structure:

{{
  "results": [
    {{
      "id": "requirement id from checklist",
      "category": "category name",
      "status": "Met" | "Partially Met" | "Not Addressed",
      "severity": "High" | "Medium" | "Low",
      "explanation": "1-2 sentence explanation of why, referencing what was or wasn't found in the document"
    }}
  ],
  "recommendations": [
    "A short, specific, actionable recommendation addressing one of the top gaps found",
    "Another recommendation",
    "..."
  ]
}}

Include an entry for every requirement in the checklist, in the same order. Be concise but specific in each explanation.
For recommendations, provide 3 to 6 of the most important, actionable next steps, prioritizing High severity gaps first. Each should be a single actionable sentence."""

    return prompt

def calculate_compliance_percentage(results):
    """Met = 1 point, Partially Met = 0.5 points, Not Addressed = 0 points."""
    if not results:
        return 0

    score_map = {"Met": 1, "Partially Met": 0.5, "Not Addressed": 0}
    total_score = sum(score_map.get(item.get("status", ""), 0) for item in results)
    percentage = (total_score / len(results)) * 100
    return round(percentage, 1)

@app.route("/")
def home():
    return "Backend is running!"

@app.route("/checklist-info", methods=["GET"])
def checklist_info():
    """Returns basic info about the checklist for the 'How it works' page."""
    categories = sorted(set(item["category"] for item in PDPA_CHECKLIST))
    return {
        "total_requirements": len(PDPA_CHECKLIST),
        "categories": categories
    }

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return "No file received", 400

    file = request.files["file"]

    if file.filename == "":
        return "No file selected", 400

    filename = file.filename
    extension = os.path.splitext(filename)[1].lower()

    if extension not in ALLOWED_EXTENSIONS:
        return f"Unsupported file type: {extension}. Please upload a PDF or DOCX file.", 400

    file.seek(0, os.SEEK_END)
    file_size_mb = file.tell() / (1024 * 1024)
    file.seek(0)

    if file_size_mb > MAX_FILE_SIZE_MB:
        return f"File too large ({file_size_mb:.1f}MB). Max size is {MAX_FILE_SIZE_MB}MB.", 400

    try:
        if extension == ".pdf":
            extracted_text = extract_text_from_pdf(file)
        elif extension == ".docx":
            extracted_text = extract_text_from_docx(file)
    except Exception as e:
        return f"Could not read this file. It may be corrupted or in an unsupported format. Error: {e}", 400

    if not extracted_text.strip():
        return "No readable text found in this document. It may be a scanned/image-based file, which isn't supported yet.", 400

    print(f"Extracted {len(extracted_text)} characters from {filename}")

    prompt = build_gap_analysis_prompt(extracted_text)

    try:
        response = client.models.generate_content(
            model="gemini-3.6-flash",
            contents=prompt
        )
    except Exception as e:
        return f"AI analysis failed: {e}", 500

    raw_text = response.text.strip()
    if raw_text.startswith("```"):
        raw_text = raw_text.strip("`")
        if raw_text.startswith("json"):
            raw_text = raw_text[4:]
        raw_text = raw_text.strip()

    try:
        analysis = json.loads(raw_text)
    except json.JSONDecodeError:
        return f"AI response could not be parsed. Raw response: {raw_text[:500]}", 500

    results = analysis.get("results", [])
    recommendations = analysis.get("recommendations", [])
    compliance_percentage = calculate_compliance_percentage(results)

    return {
        "filename": filename,
        "compliance_percentage": compliance_percentage,
        "results": results,
        "recommendations": recommendations
    }

if __name__ == "__main__":
    print("About to start Flask server...", flush=True)
    try:
        app.run(debug=True, port=5001)
    except Exception as e:
        print(f"ERROR: {e}", flush=True)
    print("Flask server has stopped.", flush=True)